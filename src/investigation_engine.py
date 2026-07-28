from docx import Document


class InvestigationEngine:

    def compare_totals(self, dashboard_total, source_total):

        difference = source_total - dashboard_total

        return {
            "dashboard_total": dashboard_total,
            "source_total": source_total,
            "difference": difference
        }

    def detect_missing_records(
        self,
        dashboard_records,
        source_records
    ):

        missing = []

        for record in source_records:
            if record not in dashboard_records:
                missing.append(record)

        return missing

    def detect_duplicates(self, records):

        seen = []
        duplicates = []

        for record in records:

            if record in seen and record not in duplicates:
                duplicates.append(record)

            seen.append(record)

        return duplicates

    def validate_date_range(
        self,
        dashboard_df,
        source_df
    ):

        dashboard_min = dashboard_df["sale_date"].min()
        dashboard_max = dashboard_df["sale_date"].max()

        source_min = source_df["sale_date"].min()
        source_max = source_df["sale_date"].max()

        issue = "No date range issue detected."

        if dashboard_max != source_max:
            issue = "Date range mismatch detected."

        return {
            "dashboard_start": dashboard_min,
            "dashboard_end": dashboard_max,
            "source_start": source_min,
            "source_end": source_max,
            "issue": issue
        }

    def validate_refresh_status(
        self,
        dashboard_df,
        source_df
    ):

        dashboard_last_date = dashboard_df["sale_date"].max()
        source_last_date = source_df["sale_date"].max()

        issue = "No refresh issue detected."

        if dashboard_last_date < source_last_date:
            issue = (
                "Dashboard refresh appears outdated. "
                "Latest source transactions are not available in dashboard."
            )

        return {
            "dashboard_last_date": dashboard_last_date,
            "source_last_date": source_last_date,
            "issue": issue
        }

    def analyze_root_cause(
        self,
        missing_records,
        duplicate_records,
        financial_impact
    ):

        analysis = []

        if len(missing_records) > 0:

            analysis.append(
                f"""
Cause 1:
{len(missing_records)} transactions are missing from dashboard data.

Affected Invoice IDs:
{missing_records}

Estimated Financial Impact:
{financial_impact}
"""
            )

        if len(duplicate_records) > 0:

            analysis.append(
                f"""
Cause 2:
{len(duplicate_records)} duplicate transaction(s) detected.

Affected Invoice IDs:
{duplicate_records}

Possible Reason:
Duplicate load during ETL process or source data issue.
"""
            )

        if len(analysis) == 0:

            analysis.append(
                """
No obvious root cause detected.
Further investigation is required.
"""
            )

        return analysis

    def calculate_impact(
        self,
        missing_records,
        duplicate_records
    ):

        return {
            "missing_count": len(missing_records),
            "duplicate_count": len(duplicate_records)
        }

    def calculate_financial_impact(
        self,
        source_df,
        missing_records
    ):

        impact = source_df[
            source_df["invoice_id"].isin(missing_records)
        ]["amount"].sum()

        return impact

    def calculate_confidence_score(
        self,
        missing_records,
        duplicate_records
    ):

        score = 100

        if len(missing_records) > 0:
            score -= 5

        if len(duplicate_records) > 0:
            score -= 5

        if score < 0:
            score = 0

        return score

    def calculate_risk_level(
        self,
        financial_impact
    ):

        if financial_impact > 5000:
            return "HIGH"

        elif financial_impact >= 1000:
            return "MEDIUM"

        return "LOW"

    def generate_summary(
        self,
        impact,
        financial_impact,
        confidence,
        date_validation,
        refresh_validation
    ):

        summary = []

        if impact["missing_count"] > 0:
            summary.append(
                f"{impact['missing_count']} missing transactions detected"
            )

        if impact["duplicate_count"] > 0:
            summary.append(
                f"{impact['duplicate_count']} duplicate transactions detected"
            )

        if date_validation["issue"] != \
                "No date range issue detected.":
            summary.append(
                "Date range mismatch detected"
            )

        if refresh_validation["issue"] != \
                "No refresh issue detected.":
            summary.append(
                "Dashboard refresh appears outdated"
            )

        return {
            "risk_level":
                self.calculate_risk_level(financial_impact),
            "financial_impact":
                financial_impact,
            "confidence":
                confidence,
            "key_findings":
                summary
        }

    def generate_report(
        self,
        total_result,
        root_cause_analysis,
        impact,
        financial_impact,
        confidence,
        date_validation
    ):

        report = f"""
====================================
INVESTIGATION REPORT
====================================

Dashboard Total      : {total_result['dashboard_total']}
Source Total         : {total_result['source_total']}
Difference           : {total_result['difference']}

Missing Count        : {impact['missing_count']}
Duplicate Count      : {impact['duplicate_count']}

Financial Impact     : {financial_impact}

Date Validation

Dashboard Range:
{date_validation['dashboard_start']} to {date_validation['dashboard_end']}

Source Range:
{date_validation['source_start']} to {date_validation['source_end']}

Issue:
{date_validation['issue']}

Detailed Root Cause Analysis:
"""

        for item in root_cause_analysis:
            report += f"\n{item}\n"

        report += f"""

Confidence Score     : {confidence}%

Recommendation:
Review missing invoices, duplicate records,
ETL processes, date filters and source data quality.

====================================
"""

        return report

    def export_report_to_docx(
        self,
        summary,
        report_text,
        output_path
    ):

        document = Document()

        document.add_heading(
            "Fabric Root Cause Agent Report",
            level=0
        )

        document.add_heading(
            "Executive Summary",
            level=1
        )

        document.add_paragraph(
            f"Risk Level: {summary['risk_level']}"
        )

        document.add_paragraph(
            f"Financial Impact: {summary['financial_impact']}"
        )

        document.add_paragraph(
            f"Confidence Score: {summary['confidence']}%"
        )

        document.add_heading(
            "Key Findings",
            level=1
        )

        for finding in summary["key_findings"]:
            document.add_paragraph(
                finding,
                style="List Bullet"
            )

        document.add_heading(
            "Detailed Investigation Report",
            level=1
        )

        document.add_paragraph(report_text)

        document.save(output_path)

        return output_path
